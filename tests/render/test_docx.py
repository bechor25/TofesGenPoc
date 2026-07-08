from docx import Document

from doc2tests.contracts.enums import FieldType, SourceKind
from doc2tests.contracts.records import Record, Value
from doc2tests.contracts.template import CanonicalTemplate, DocSource, Field
from doc2tests.render.docx import render_docx


def _tmpl():
    return CanonicalTemplate(
        doc_type="bank-form", source=DocSource(kind=SourceKind.image),
        fields=[Field(id="pid", label="מספר זהות", type=FieldType.israeli_id)],
    )


def test_docx_written_and_contains_values(tmp_path):
    out = tmp_path / "r.docx"
    rec = Record(index=0, test_class="equivalence", expected_valid=True,
                 values={"pid": Value(field_id="pid", value="123456782")})
    render_docx(_tmpl(), rec, str(out))
    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "bank-form" in text
    assert "מספר זהות" in table_text
    assert "123456782" in table_text
