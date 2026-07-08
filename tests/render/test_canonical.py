import json

from docx import Document

from doc2tests.contracts.enums import FieldType, SourceKind, TestClass
from doc2tests.contracts.records import Record, Value
from doc2tests.contracts.template import CanonicalTemplate, DocSource, Field
from doc2tests.render.canonical import (
    dataset_csv,
    dataset_json,
    fill_docx_from_template,
    records_from_rows,
    render_blank_docx,
    render_blank_html,
    template_json,
)


def _tmpl():
    return CanonicalTemplate(
        doc_type="bank-form", source=DocSource(kind=SourceKind.image),
        fields=[
            Field(id="pid", label="מספר זהות", type=FieldType.israeli_id),
            Field(id="entry_date", label="תאריך כניסה", type=FieldType.date),
        ],
    )


def test_template_json_roundtrips():
    js = template_json(_tmpl())
    reloaded = CanonicalTemplate.model_validate_json(js)
    assert reloaded.doc_type == "bank-form"
    assert json.loads(js)["fields"][0]["id"] == "pid"


def test_blank_html_shows_placeholders_not_values():
    html = render_blank_html(_tmpl())
    assert "{{ pid }}" in html
    assert "{{ entry_date }}" in html
    assert 'dir="rtl"' in html
    assert "מספר זהות" in html


def test_blank_docx_is_a_fillable_docxtpl_template(tmp_path):
    blank = tmp_path / "template.docx"
    render_blank_docx(_tmpl(), str(blank))
    assert blank.exists()
    # placeholders present as text
    doc = Document(str(blank))
    text = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "{{ pid }}" in text

    # and docxtpl can fill it with arbitrary data
    out = tmp_path / "filled.docx"
    fill_docx_from_template(str(blank), {"pid": "123456782", "entry_date": "31.10.21"},
                            str(out))
    filled = Document(str(out))
    filled_text = " ".join(c.text for t in filled.tables for row in t.rows for c in row.cells)
    assert "123456782" in filled_text
    assert "{{ pid }}" not in filled_text


def test_records_from_rows_validates_user_data():
    recs = records_from_rows(_tmpl(), [
        {"pid": "123456782", "entry_date": "31.10.21"},   # valid
        {"pid": "123456789", "entry_date": "bad"},         # invalid id + date
    ])
    assert recs[0].expected_valid is True
    assert recs[1].expected_valid is False
    assert recs[1].values["pid"].valid is False


def test_records_from_rows_accepts_label_keys():
    recs = records_from_rows(_tmpl(), [{"מספר זהות": "123456782", "תאריך כניסה": "01.01.2020"}])
    assert recs[0].values["pid"].value == "123456782"


def test_dataset_json_and_csv():
    pop = [Record(index=0, test_class=TestClass.equivalence, expected_valid=True,
                  values={"pid": Value(field_id="pid", value="123456782"),
                          "entry_date": Value(field_id="entry_date", value="01.01.2020")})]
    js = json.loads(dataset_json(pop))
    assert js[0]["values"]["pid"] == "123456782"
    csv_text = dataset_csv(_tmpl(), pop)
    assert "pid" in csv_text.splitlines()[0]
    assert "123456782" in csv_text
