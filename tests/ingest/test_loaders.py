import fitz
import pytest
from docx import Document

from doc2tests.ingest.loaders import detect_kind, load_images, read_docx_text


def test_detect_kind_by_extension():
    assert detect_kind("a.jpg") == "image"
    assert detect_kind("a.PNG") == "image"
    assert detect_kind("a.pdf") == "pdf"
    assert detect_kind("a.docx") == "docx"


def test_detect_kind_rejects_unknown():
    with pytest.raises(ValueError):
        detect_kind("a.txt")


def test_load_images_passthrough_for_image(tmp_path):
    p = tmp_path / "x.png"
    p.write_bytes(b"\x89PNG\r\n")
    assert load_images(str(p)) == [b"\x89PNG\r\n"]


def test_load_images_renders_pdf_pages(tmp_path):
    p = tmp_path / "x.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.new_page()
    doc.save(str(p))
    doc.close()
    images = load_images(str(p))
    assert len(images) == 2
    assert images[0].startswith(b"\x89PNG")  # rendered to PNG


def test_read_docx_text_extracts_paragraphs_and_tables(tmp_path):
    p = tmp_path / "d.docx"
    doc = Document()
    doc.add_paragraph("בקשה להעברת תעודת זכאות")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "מספר זהות"
    table.rows[0].cells[1].text = "123456782"
    doc.save(str(p))
    text = read_docx_text(str(p))
    assert "בקשה להעברת תעודת זכאות" in text
    assert "מספר זהות" in text
    assert "123456782" in text
