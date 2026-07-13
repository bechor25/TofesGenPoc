from pathlib import Path

import pytest

from doc2tests.ingest.rasterize import rasterize, soffice_path

FIXTURES = Path(__file__).parent.parent / "fixtures"


def _first_image_fixture() -> Path:
    for p in FIXTURES.iterdir():
        if p.suffix.lower() in {".jpeg", ".jpg", ".png"}:
            return p
    raise AssertionError("no image fixture")


def test_image_passthrough_returns_png_bytes():
    src = _first_image_fixture()
    pages = rasterize(str(src))
    assert len(pages) >= 1
    assert pages[0][:4] == b"\x89PNG"  # normalized to PNG


def test_unknown_kind_raises():
    with pytest.raises(ValueError):
        rasterize("form.txt")


@pytest.mark.skipif(soffice_path() is None, reason="LibreOffice not installed")
def test_docx_renders_at_least_one_page(tmp_path):
    from docx import Document
    docx_path = tmp_path / "f.docx"
    d = Document()
    d.add_paragraph("שם: דנה כהן")
    d.add_paragraph("ת.ז: 123456782")
    d.save(docx_path)
    pages = rasterize(str(docx_path))
    assert len(pages) >= 1
    assert pages[0][:4] == b"\x89PNG"
