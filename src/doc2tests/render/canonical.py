"""Template-as-document artifacts: the blank template (placeholders, no values),
canonical JSON export, dataset export, and filling the template with *arbitrary*
user-supplied records — not only the generated population."""
from __future__ import annotations

import csv
import io
import json
from html import escape

from docx import Document
from docxtpl import DocxTemplate

from doc2tests.contracts.enums import TestClass
from doc2tests.contracts.records import Record, Value
from doc2tests.contracts.template import CanonicalTemplate
from doc2tests.render.docxutil import set_rtl
from doc2tests.render.style import html_shell
from doc2tests.validators import validate


def template_json(template: CanonicalTemplate) -> str:
    """The canonical template — the single source of truth — as JSON."""
    return template.model_dump_json(indent=2)


def render_blank_html(template: CanonicalTemplate) -> str:
    """The document AS a template: every value cell shows its placeholder token,
    no data filled in."""
    rows = []
    for f in template.fields:
        rows.append(
            f'<tr><td class="label">{escape(f.label)}'
            f'<span class="tag">{escape(f.type.value)}</span></td>'
            f'<td><span class="ph">{escape(f.placeholder)}</span></td></tr>'
        )
    meta = '<span class="badge badge--ok">טמפלייט</span>' \
           '<span class="badge">ללא ערכים</span>'
    foot = f"doc_type: {escape(template.doc_type)} · {len(template.fields)} placeholders"
    return html_shell(f"{template.doc_type} — טמפלייט", meta, "\n".join(rows), foot)


def render_blank_docx(template: CanonicalTemplate, path: str) -> None:
    """A reusable DOCX template with ``{{ field_id }}`` placeholders — fillable
    later by docxtpl (see :func:`fill_docx_from_template`)."""
    doc = Document()
    heading = doc.add_heading(f"{template.doc_type} — טמפלייט", level=1)
    set_rtl(heading)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = 2  # right
    hdr = table.rows[0].cells
    hdr[0].text = "שדה"
    hdr[1].text = "ערך"
    for cell in hdr:
        set_rtl(cell.paragraphs[0])
    for f in template.fields:
        cells = table.add_row().cells
        cells[0].text = f.label
        cells[1].text = f"{{{{ {f.id} }}}}"  # docxtpl placeholder
        for cell in cells:
            set_rtl(cell.paragraphs[0])
    doc.save(path)


def fill_docx_from_template(blank_path: str, values: dict[str, str], out_path: str) -> None:
    """Fill a blank DOCX template (from :func:`render_blank_docx`) with any values."""
    tpl = DocxTemplate(blank_path)
    tpl.render(values)
    tpl.save(out_path)


def records_from_rows(
    template: CanonicalTemplate, rows: list[dict[str, str]]
) -> list[Record]:
    """Turn user-supplied rows (keyed by field id, or by label) into validated
    Records that can be rendered exactly like the generated population."""
    records: list[Record] = []
    for i, row in enumerate(rows):
        values: dict[str, Value] = {}
        all_valid = True
        for f in template.fields:
            raw = row.get(f.id)
            if raw is None:
                raw = row.get(f.label, "")
            raw = str(raw)
            ok = validate(f.type, raw) if raw else not f.constraints.required
            all_valid = all_valid and ok
            values[f.id] = Value(field_id=f.id, value=raw, valid=ok)
        records.append(Record(index=i, test_class=TestClass.equivalence,
                              expected_valid=all_valid, values=values))
    return records


def dataset_json(population: list[Record]) -> str:
    data = [
        {"index": r.index, "test_class": str(r.test_class),
         "expected_valid": r.expected_valid, "violates": r.violates,
         "values": {fid: v.value for fid, v in r.values.items()}}
        for r in population
    ]
    return json.dumps(data, ensure_ascii=False, indent=2)


def dataset_csv(template: CanonicalTemplate, population: list[Record]) -> str:
    field_ids = [f.id for f in template.fields]
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["index", "test_class", "expected_valid", "violates", *field_ids])
    for r in population:
        writer.writerow([
            r.index, str(r.test_class), r.expected_valid, r.violates or "",
            *[r.values[fid].value if fid in r.values else "" for fid in field_ids],
        ])
    return buf.getvalue()
