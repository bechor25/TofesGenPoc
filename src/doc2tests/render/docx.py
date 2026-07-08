from __future__ import annotations

from docx import Document

from doc2tests.contracts.records import Record
from doc2tests.contracts.template import CanonicalTemplate
from doc2tests.render.docxutil import set_rtl


def render_docx(template: CanonicalTemplate, record: Record, path: str) -> None:
    doc = Document()
    heading = doc.add_heading(template.doc_type, level=1)
    set_rtl(heading)

    meta = doc.add_paragraph(
        f"רשומה #{record.index} · {record.test_class}"
        + ("" if record.expected_valid else f" · expected INVALID ({record.violates})")
    )
    set_rtl(meta)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = 2  # right
    hdr = table.rows[0].cells
    hdr[0].text = "שדה"
    hdr[1].text = "ערך"
    set_rtl(hdr[0].paragraphs[0])
    set_rtl(hdr[1].paragraphs[0])
    for f in template.fields:
        v = record.values.get(f.id)
        cells = table.add_row().cells
        cells[0].text = f.label
        cells[1].text = v.value if v else ""
        set_rtl(cells[0].paragraphs[0])
        set_rtl(cells[1].paragraphs[0])
    doc.save(path)
