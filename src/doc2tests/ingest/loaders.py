"""Format-aware loading: turn any supported input (image / PDF / Word) into
the representation the extractor needs — images for vision, text for Word."""
from __future__ import annotations

from pathlib import Path
from typing import Literal

Kind = Literal["image", "pdf", "docx"]

_IMAGE_EXT = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}
_MAX_PDF_PAGES = 3


def detect_kind(path: str) -> Kind:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in (".docx", ".doc"):
        return "docx"
    if ext in _IMAGE_EXT:
        return "image"
    raise ValueError(f"unsupported file type: {ext or path}")


def _render_pdf_to_images(path: str, dpi: int = 200) -> list[bytes]:
    import fitz  # PyMuPDF

    images: list[bytes] = []
    with fitz.open(path) as doc:
        for page in doc[:_MAX_PDF_PAGES]:
            pix = page.get_pixmap(dpi=dpi)
            images.append(pix.tobytes("png"))
    return images


def load_images(path: str) -> list[bytes]:
    """Return page images (PNG/JPEG bytes) for an image or PDF input."""
    kind = detect_kind(path)
    if kind == "image":
        return [Path(path).read_bytes()]
    if kind == "pdf":
        return _render_pdf_to_images(path)
    raise ValueError(f"{path} is not an image/pdf")


def read_docx_text(path: str) -> str:
    """Flatten a Word document (paragraphs + table cells) to plain text."""
    from docx import Document

    doc = Document(path)
    lines: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)
